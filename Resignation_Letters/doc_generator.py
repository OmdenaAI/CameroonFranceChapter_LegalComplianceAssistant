import itertools
import os
from pathlib import Path
import random
import time

from docx import Document
from docx.shared import Cm
import faker

SIGS_DIR = Path("signatures")
OUTPUT_DIR = Path("generated")
SIGNATURE_KEY = "[SIGNATURE]"
POSITIONS = (
    "Software Engineer",
    "DevOps",
    "Cloud Architect",
    "Site Reliability Engineer",
    "Data Engineer",
    "Data Analyst",
)

SENIORITY = ("Junior", "Mid", "Senior", "Staff")


def generate_fake_data(num: int) -> list[dict[str, str]]:
    fake = faker.Faker()
    roles: list[str] = []
    for seniority, position in itertools.product(SENIORITY, POSITIONS):
        roles.append(f"{seniority} {position}")

    data: list[dict[str, str]] = []
    for _ in range(num):
        data.append(
            {
                "[NAME]": fake.name(),
                "[ADDRESS]": fake.address(),
                "[CURRENT DATE]": fake.date(),
                "[RECIPIENT NAME]": fake.name(),
                "[RECIPIENT ADDRESS]": fake.address(),
                "[POSITION NAME]": random.choice(roles),
                "[COMPANY NAME]": fake.company(),
                "[NUMBER OF WEEKS NOTICE]": str(random.randrange(2, 13, 1)),
            }
        )

    return data


def generate_doc(
    template: Path, data: dict[str, str], signatures: list[Path], output: Path
) -> None:
    doc = Document(str(template))

    for paragraph in doc.paragraphs:
        for key, value in data.items():
            paragraph.text = paragraph.text.replace(key, value)
        if SIGNATURE_KEY in paragraph.text:
            paragraph.text = paragraph.text.replace(SIGNATURE_KEY, "")
            run = paragraph.add_run()
            sig_path = random.choice(signatures)
            sig_width = random.randrange(200, 301, 1) / 100
            run.add_picture(str(sig_path), Cm(sig_width))

    doc.save(output)


def get_signatures(sigs_dir: Path = SIGS_DIR) -> list[Path]:
    signatures_filepaths: list[Path] = []
    for root, _, filenames in os.walk(sigs_dir):
        for filename in filenames:
            if not filename.endswith(".png"):
                continue
            signatures_filepaths.append(Path(os.path.join(root, filename)))

    return signatures_filepaths


def generate_resignation_letters(num: int, output_dir: Path = OUTPUT_DIR) -> None:
    template = Path("./Resignation-letter-Template.docx")
    signatures = get_signatures()
    data_dicts: list[dict[str, str]] = generate_fake_data(num)
    for n, data in enumerate(data_dicts):
        generate_doc(
            template, data, signatures, output_dir / f"resignation_letter_{n}.docx"
        )


if __name__ == "__main__":
    start = time.perf_counter()
    generate_resignation_letters(20)
    end = time.perf_counter()
    print(f"Documents generated in {end - start:.2f} seconds")
