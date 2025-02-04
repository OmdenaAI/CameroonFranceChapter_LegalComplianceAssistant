from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from faker import Faker
import random
import os

fake = Faker()

output_folder = "generated_bills"
os.makedirs(output_folder, exist_ok=True)

def generate_random_phone_number():
    # Generate a random 3-digit, 3-digit, and 4-digit phone number
    area_code = random.randint(100, 999)
    central_office_code = random.randint(100, 999)
    subscriber_number = random.randint(1000, 9999)
    
    return f"{area_code}-{central_office_code}-{subscriber_number}"


def get_random_signature():
    signature_folder = "signatures"
    signature_files = [f for f in os.listdir(signature_folder) if f.endswith('.png') or f.endswith('.jpg')]
    
    if not signature_files:
        print("No signature images found in the 'signatures' folder!")
        return
    
    signature_image = random.choice(signature_files)
    
    signature_path = os.path.join(signature_folder, signature_image)
    
    return signature_path

def create_fake_hospital_bill(file_name):
    document = Document()

    # logo
    document.add_picture("logo.png", width=Inches(2))  

    hospital_names = [
    "UVM Medical Center",
    "St. Mary's Health Hospital",
    "Green Valley Medical Center",
    "Riverbend Hospital",
    "Springfield General Hospital",
    "Blue Cross Health Institute",
    "Summit Medical Center",
    "Hopewell Healthcare",
    "UnityPoint Hospital",
    "Northern Lights Medical Center"
]

    # fake patient data
    patient_name = fake.name()
    dob = fake.date_of_birth()
    gender = random.choice(["Male", "Female", "Non-binary", "Other"])
    address = fake.address().replace("\n", ", ")
    account_number = fake.random_number(digits=8)
    insurance = random.choice(["Aetna", "Cigna", "UnitedHealthcare", 
                               "Blue Cross Blue Shield", "Humana", "Ambetter", "Centene", "Anthem", 
                               "Molina Healthcare", "Kaiser Permanente", "Bright Health", "Elevance Health","Health Net" ])
    statement_date = fake.date_this_decade()
    due_date = fake.date_between_dates(statement_date, statement_date.replace(year=statement_date.year + 1))

    # hospital info
    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hospital_name = random.choice(hospital_names)
    title_run = title.add_run(hospital_name)
    title_run.bold = True
    title_run.font.size = Pt(14)

    locations = [
    ("New York", "NY"),
    ("Los Angeles", "CA"),
    ("Chicago", "IL"),
    ("Houston", "TX"),
    ("Phoenix", "AZ"),
    ("Philadelphia", "PA"),
    ("San Antonio", "TX"),
    ("San Diego", "CA"),
    ("Dallas", "TX"),
    ("San Jose", "CA"),
    ("Denver", "CO"),
    ("Seattle", "WA"),
    ("Miami", "FL"),
    ("Boston", "MA"),
    ("Atlanta", "GA")
    ]

    department_name = random.choice(["Payments Department", "Billing Department", "Patient Services", "Healthcare Payments", "Financial Services"])

    po_box = f"PO BOX {random.randint(1000, 9999)}"
    city, state = random.choice(locations)
    zip_code = random.randint(10000, 99999)
    document.add_paragraph(f"{po_box}\n{city}, {state} {zip_code}\nRETURN SERVICE REQUESTED\n")

    document.add_paragraph("\n")

    # Patient Information
    patient_info = document.add_paragraph()
    patient_info.add_run(f"PATIENT NAME:  ").bold = True
    patient_info.add_run(f"{patient_name}\n")
    patient_info.add_run("DATE OF BIRTH:  ").bold = True
    patient_info.add_run(f"{dob}\n")
    patient_info.add_run("GENDER:  ").bold = True
    patient_info.add_run(f"{gender}\n")
    patient_info.add_run("ADDRESS:  ").bold = True
    patient_info.add_run(f"{address}\n")
    patient_info.add_run("ACCOUNT #:  ").bold = True
    patient_info.add_run(f"{account_number}\n")
    patient_info.add_run("INSURANCE INFORMATION ON FILE:  ").bold = True
    patient_info.add_run(f"{insurance}\n")

    document.add_paragraph("\n")

    # Contact Details
    contact_info = document.add_paragraph()
    contact_info.add_run("BILLING QUESTIONS:\n").bold = True
    tel_number = generate_random_phone_number()
    toll_free_number = generate_random_phone_number()
    contact_info.add_run(f"TEL: {tel_number}  |  TOLL-FREE: {toll_free_number}\n")

    document.add_paragraph("\n")

    # Invoice Header
    invoice_details = document.add_paragraph()
    invoice_details.add_run("STATEMENT DATE:  ").bold = True
    invoice_details.add_run(f"{statement_date.strftime('%m/%d/%Y')}\n")
    invoice_details.add_run("DUE DATE:  ").bold = True
    invoice_details.add_run(f"{due_date.strftime('%m/%d/%Y')}\n")

    document.add_paragraph("\n")

    # Billing Table
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "ADMIT DATE OF SERVICE"
    hdr_cells[1].text = "VISIT NUMBER / PROCEDURE CODE"
    hdr_cells[2].text = "TYPE OF SERVICE"
    hdr_cells[3].text = "BILLED CHARGES"
    hdr_cells[4].text = "PAID BY PLAN / ADJUSTMENT"
    hdr_cells[5].text = "PATIENT PAYMENT DUE"

    # generate fake bill
    billing_data = []
    total_due = 0

    for _ in range(random.randint(1, 5)):  # 1-5 random medical services
        service_date = fake.date_this_decade()
        visit_number = fake.random_number(digits=8)
        procedure_code = fake.random_int(min=100, max=999)
        service_type = random.choice([
                                "Outpatient Consultation", "Emergency Room Visit", "Routine Checkup",
                                "Lab Tests", "X-Ray", "CT Scan", "MRI Scan", "Ultrasound",
                                "Physical Therapy", "Blood Work", "Vaccination", "ECG Test",
                                "Skin Biopsy", "Endoscopy", "Colonoscopy", "Dialysis", 
                                "Cardiology Consultation", "Neurology Consultation", "Pulmonology Exam"
                            ])

        billed_charges = round(random.uniform(50, 300), 2)
        paid_by_plan = round(random.uniform(10, billed_charges - 10), 2)
        patient_due = round(billed_charges - paid_by_plan, 2)
        total_due += patient_due

        billing_data.append((service_date.strftime('%m/%d/%Y'), f"{visit_number}\n{procedure_code}", service_type, f"${billed_charges}", f"-${paid_by_plan}", f"${patient_due}"))

    for row in billing_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row):
            row_cells[i].text = text

    document.add_paragraph("\n")

    # Amount Due
    amount_due = document.add_paragraph()
    amount_due.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    amount_due_run = amount_due.add_run(f"ACCOUNT TOTAL DUE FROM YOU:  ${total_due:.2f}")
    amount_due_run.bold = True
    amount_due_run.font.size = Pt(12)

    document.add_paragraph("\n")

    # Payment Instructions
    payment_instructions = document.add_paragraph()
    payment_instructions.add_run("IMPORTANT MESSAGE:\n").bold = True
    payment_instructions.add_run(f"Thank you for selecting {hospital_name} as your healthcare provider. Please pay the amount due by the due date shown on your statement.\n")

    document.add_paragraph("\n")

    # Payment Slip
    payment_slip = document.add_paragraph()
    payment_slip.add_run("MAKE CHECKS PAYABLE TO:\n").bold = True

   
    document.add_paragraph(f"{hospital_name}\n{department_name}\n{po_box}\n{city}, {state} {zip_code}\n")
    document.add_paragraph("\n")

    # Signature
    signature_section = document.add_paragraph()
    signature_section.add_run("Authorized Signature: ___________________________\n").bold = True
    signature_section.add_run("Date: ______________")

    # Add Signature Image
    signature_path = get_random_signature()
    document.add_picture(signature_path, width=Inches(2)) 

    document.add_paragraph("\n")

    # Footer
    footer = document.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer.add_run("Page 2 of 2")
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(0, 128, 0)  # Green color

    # Save Document
    document_path = os.path.join(output_folder, file_name)
    document.save(document_path)
    print(f"Hospital Bill saved as '{document_path}'")

# Generate 5 Fake Bills
for i in range(5):
    file_name = f"hospital_bill_{i+1}.docx"
    create_fake_hospital_bill(file_name)

print("All bills generated successfully!")
