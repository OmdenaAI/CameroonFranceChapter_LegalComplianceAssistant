from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from faker import Faker
import random
import os
faker_object = Faker()
drive_link = r'<drivelink>'
# Function to define the details
def get_divorce_details():
    return {
        "husband_name": faker_object.name_male(),
        "wife_name": faker_object.name_female(),
        "marriage_date": faker_object.date_this_year(),
        "marriage_location": faker_object.address(),
        "husband_address": faker_object.address(),
        "wife_address": faker_object.address(),
        "property_address": faker_object.address(),
        "child_support_amount": f"${str(random.randint(100,1000))}",
        "spousal_support_amount": f"${str(random.randint(100,1000))}",
        "spousal_support_duration": f"{str(random.randint(1,10))} years",
        "life_insurance_amount": f"${str(random.randint(10000,1000000))}",
        "payment_method": "bank transfer",
        "start_date": faker_object.date_this_year(),
        "witness_1_name": faker_object.name(),
        "witness_2_name": faker_object.name(),
        "notary_name": faker_object.company(),
        "state": faker_object.state(),
        "county": faker_object.country()
    }

# Function to generate the divorce agreement

def get_signature_path():
    folder_choice = random.choice(os.listdir(f"{drive_link}/signature/Dataset_Signature_Final/Dataset"))
    subfolder = random.choice(['forge','real'])
    sign_choice = random.choice(os.listdir(f"{drive_link}/signature/Dataset_Signature_Final/Dataset/{folder_choice}/{subfolder}"))
    signature_path = f"{drive_link}/signature/Dataset_Signature_Final/Dataset/{folder_choice}/{subfolder}/{sign_choice}"
    return signature_path
def generate_divorce_agreement(details, i):
    # Create a new Word document
    doc = Document()

    # Add a title
    title = doc.add_heading('Divorce Agreement', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add an introduction paragraph
    intro = doc.add_paragraph()
    intro.add_run(f'This Divorce Agreement ("Agreement") is made and entered into on this ___ day of __________, 20__, by and between ').bold = True
    intro.add_run(f'{details["husband_name"]}, residing at {details["husband_address"]}, and {details["wife_name"]}, residing at {details["wife_address"]}. '
                  f'The parties have agreed to dissolve their marriage and settle all matters related to the division of assets, child custody, spousal support, and other obligations as outlined below.')

    # Section 1: Recitals
    doc.add_heading('1. Recitals', level=2)
    recitals = doc.add_paragraph()
    recitals.add_run(f'WHEREAS, the parties were married on {details["marriage_date"]} in {details["marriage_location"]};\n')
    recitals.add_run('WHEREAS, the parties have decided to dissolve their marriage amicably;\n')
    recitals.add_run('WHEREAS, the parties wish to resolve all issues arising from their marriage, including but not limited to property division, child custody, child support, spousal support, and other matters;\n')
    recitals.add_run('NOW, THEREFORE, in consideration of the mutual promises and covenants contained herein, the parties agree as follows:')

    # Section 2: Division of Assets and Debts
    doc.add_heading('2. Division of Assets and Debts', level=2)
    assets = doc.add_paragraph()
    assets.add_run(f'The parties agree to divide their marital assets and debts as follows:\n')
    assets.add_run(f'- Real Estate: The marital home located at {details["property_address"]} shall be transferred to {details["husband_name"]}. '
                   f'Any outstanding mortgage or liens shall remain the responsibility of {details["husband_name"]}.\n')
    assets.add_run('- Bank Accounts: All joint bank accounts shall be closed, and the funds shall be divided equally between the parties.\n')
    assets.add_run('- Retirement Accounts: Each party shall retain their respective retirement accounts (e.g., 401(k), IRA). Any joint retirement accounts shall be divided according to a Qualified Domestic Relations Order (QDRO).\n')
    assets.add_run('- Vehicles: Each party shall retain ownership of the vehicle(s) titled in their name. Any loans associated with the vehicles shall remain the responsibility of the titled owner.\n')
    assets.add_run('- Personal Property: Each party shall retain their personal belongings, including furniture, electronics, and other household items as mutually agreed upon.\n')
    assets.add_run('- Debts: All joint debts incurred during the marriage shall be divided equally unless otherwise specified herein.')

    # Section 3: Child Custody and Support
    doc.add_heading('3. Child Custody and Support', level=2)
    custody = doc.add_paragraph()
    custody.add_run('The parties agree to the following custody and support arrangement:\n')
    custody.add_run('- Legal Custody: Joint legal custody shall be granted to both parties, meaning both parents shall share decision-making authority regarding the children’s education, healthcare, and general welfare.\n')
    custody.add_run('- Physical Custody: The children shall reside primarily with [Parent Name], with visitation rights granted to [Other Parent Name] as follows:\n')
    custody.add_run('  - Weekends: Every other weekend from Friday at 6:00 PM to Sunday at 6:00 PM.\n')
    custody.add_run('  - Holidays: The children shall alternate holidays between the parents as follows: Thanksgiving with [Parent A] in even years, Christmas with [Parent B] in odd years, etc.\n')
    custody.add_run(f'- Child Support: {details["husband_name"]} shall pay {details["child_support_amount"]} per month in child support to {details["wife_name"]}. '
                    f'Payments shall be made via {details["payment_method"]} and are due on the 1st of each month.\n')
    custody.add_run('- Medical Expenses: Both parties shall equally share any uninsured medical, dental, or educational expenses for the children.')

    # Section 4: Spousal Support (Alimony)
    doc.add_heading('4. Spousal Support (Alimony)', level=2)
    spousal_support = doc.add_paragraph()
    spousal_support.add_run(f'The parties agree to the following spousal support arrangement:\n')
    spousal_support.add_run(f'- {details["husband_name"]} shall pay {details["spousal_support_amount"]} per month in spousal support to {details["wife_name"]} for a period of {details["spousal_support_duration"]}, beginning on {details["start_date"]}.\n')
    spousal_support.add_run('- The obligation to pay spousal support shall terminate upon the death of either party, the remarriage of the receiving party, or the expiration of the agreed-upon duration, whichever occurs first.\n')

    # Section 5: Health Insurance and Benefits
    doc.add_heading('5. Health Insurance and Benefits', level=2)
    insurance = doc.add_paragraph()
    insurance.add_run('The parties agree to the following regarding health insurance and benefits:\n')
    insurance.add_run(f'- Health Insurance: {details["husband_name"]} shall maintain health insurance coverage for the children until they reach the age of majority or graduate from high school, whichever occurs later.\n')
    insurance.add_run(f'- Life Insurance: {details["husband_name"]} shall maintain a life insurance policy with a face value of {details["life_insurance_amount"]}, naming {details["wife_name"]} as the beneficiary to secure child support and/or spousal support obligations.\n')

    # Section 6: Tax Matters
    doc.add_heading('6. Tax Matters', level=2)
    taxes = doc.add_paragraph()
    taxes.add_run('The parties agree to the following regarding tax matters:\n')
    taxes.add_run('- Dependency Exemptions: The parties shall alternate claiming the children as dependents for tax purposes, with [Parent A] claiming them in even years and [Parent B] in odd years.\n')
    taxes.add_run('- Filing Status: Each party shall file their taxes separately unless otherwise agreed upon in writing.\n')

    # Section 7: Dispute Resolution
    doc.add_heading('7. Dispute Resolution', level=2)
    dispute = doc.add_paragraph()
    dispute.add_run('In the event of any disputes arising under this Agreement, the parties agree to attempt to resolve such disputes through mediation before pursuing litigation. If mediation fails, the matter may be resolved through arbitration or court proceedings as permitted by law.')

    # Section 8: General Provisions
    doc.add_heading('8. General Provisions', level=2)
    general = doc.add_paragraph()
    general.add_run('The parties agree to the following general provisions:\n')
    general.add_run('- Entire Agreement: This Agreement constitutes the entire understanding between the parties and supersedes all prior agreements, whether written or oral.\n')
    general.add_run('- Modifications: Any modifications to this Agreement must be in writing and signed by both parties.\n')
    general.add_run(f'- Governing Law: This Agreement shall be governed by and construed in accordance with the laws of the State of {details["state"]}.\n')

    # Section 9: Signatures
    doc.add_heading('9. Signatures', level=2)

    # Husband's Signature
    doc.add_paragraph(f'{details["husband_name"]}\'s Signature:')
    signature_path = get_signature_path()
    doc.add_picture(signature_path, width=Inches(1.5))  # Replace with actual file path
    doc.add_paragraph('Date: ___/___/20__\n')

    # Wife's Signature
    doc.add_paragraph(f'{details["wife_name"]}\'s Signature:')
    signature_path = get_signature_path()
    doc.add_picture(signature_path, width=Inches(1.5))  # Replace with actual file path
    doc.add_paragraph('Date: ___/___/20__\n')

    # Witness 1 Signature
    doc.add_paragraph(f'Witness 1: {details["witness_1_name"]}')
    signature_path = get_signature_path()
    doc.add_picture(signature_path, width=Inches(1.5))  # Replace with actual file path
    doc.add_paragraph('Date: ___/___/20__\n')

    # Witness 2 Signature
    doc.add_paragraph(f'Witness 2: {details["witness_2_name"]}')
    signature_path = get_signature_path()
    doc.add_picture(signature_path, width=Inches(1.5))  # Replace with actual file path
    doc.add_paragraph('Date: ___/___/20__\n')

    # Notary Section
    doc.add_heading('Notary Acknowledgment', level=2)
    notary_text = doc.add_paragraph()
    notary_text.add_run(f'State of {details["state"]}\nCounty of {details["county"]}\n\n')
    notary_text.add_run('On this ___ day of __________, 20__, before me, the undersigned notary public, personally appeared ')
    notary_text.add_run(f'{details["husband_name"]} and {details["wife_name"]}, known to me (or satisfactorily proven) to be the persons whose names are subscribed to this instrument, and acknowledged that they executed it for the purposes therein contained.\n\n')
    notary_text.add_run('___________________________\n')
    notary_text.add_run(f'{details["notary_name"]}, Notary Public\n')
    notary_text.add_run('My Commission Expires: ___/___/20__\n')
    notary_seal = r'D:\Raghu Studies\omdena\CameroonFranceChapter_LegalComplianceAssistant\data\notary.jpg'
    doc.add_picture(notary_seal, width=Inches(1.5))  # Replace with actual file path

    # Save the document
    doc.save(rf'D:\Raghu Studies\omdena\CameroonFranceChapter_LegalComplianceAssistant\data\divorce\Detailed_Divorce_Agreement_{i}.docx')
    print("Detailed Divorce Agreement generated successfully!")

# Main execution
if __name__ == "__main__":
    for i in range(60):
        details = get_divorce_details()
        generate_divorce_agreement(details, i)
