import os
import random
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont
from faker import Faker
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

fake = Faker()

# --- Utility Functions ---

def get_random_font(font_folder):
    fonts = [os.path.join(font_folder, f) for f in os.listdir(font_folder)
             if f.lower().endswith('.ttf') or f.lower().endswith('.otf')]
    if not fonts:
        raise ValueError(f"No font files found in {font_folder}. Please add .ttf or .otf font files.")
    return random.choice(fonts)

def generate_signature_image(name, output_path, font_folder, width=600, height=100, font_size=48, text_color=(0, 0, 0)):
    font_path = get_random_font(font_folder)
    print(f"Using font: {font_path}")

    temp_image = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(temp_image)

    try:
        font = ImageFont.truetype(font_path, font_size)
    except IOError:
        print(f"Font file '{font_path}' not found. Using default font.")
        font = ImageFont.load_default()

    bbox = draw.textbbox((0, 0), name, font=font)
    text_width = bbox[2] - bbox[0]
    text_height = bbox[3] - bbox[1]

    canvas_width = int(text_width * 1.2)
    canvas_height = int(text_height * 1.2)
    image = Image.new("RGBA", (canvas_width, canvas_height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)
    text_x = (canvas_width - text_width) // 2
    text_y = (canvas_height - text_height) // 2
    draw.text((text_x, text_y), name, font=font, fill=text_color)

    angle = random.uniform(-5, 5)
    rotated_image = image.rotate(angle, expand=True)
    rotated_image.save(output_path)

def add_signature_above(doc, name, output_dir, fonts_folder, role):
    """
    Adds a signature image directly above the given role's signature line.
    """
    # Generate and save the signature image
    signature_path = os.path.join(output_dir, f"signature_{name.replace(' ', '_')}.png")
    generate_signature_image(name, signature_path, fonts_folder)

    # Locate the signature line and insert the image above it
    signature_line_text = f"{role}'s Signature: __________________________"
    for paragraph in doc.paragraphs:
        if signature_line_text in paragraph.text:
            # Insert a new paragraph with the image above the signature line
            img_paragraph = paragraph.insert_paragraph_before()
            img_run = img_paragraph.add_run()
            img_run.add_picture(signature_path, width=Inches(3.5))  # Adjust width if needed
            img_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            img_paragraph.paragraph_format.space_after = Pt(0)
            img_paragraph.paragraph_format.space_before = Pt(0)
            break

def create_lease_agreement(num_docs=1, output_dir="Lease_Agreements", fonts_folder="fonts"):
    os.makedirs(output_dir, exist_ok=True)

    for _ in range(num_docs):
        data = {
            "landlord_name": fake.name(),
            "landlord_address": fake.address().replace("\n", ", "),
            "tenant_name": fake.name(),
            "tenant_address": fake.address().replace("\n", ", "),
            "property_address": fake.address().replace("\n", ", "),
            "residence_type": random.choice(["Apartment", "House", "Condo", "Townhouse"]),
            "bedrooms": random.randint(1, 5),
            "bathrooms": random.randint(1, 3),
            "start_date": fake.date_between(start_date="today", end_date="+30d").strftime("%B %d, %Y"),
            "end_date": fake.date_between(start_date="+1y", end_date="+2y").strftime("%B %d, %Y"),
            "rent_amount": random.randint(1000, 3000),
            "security_deposit": random.randint(1500, 4500),
            "agreement_date": fake.date_this_month().strftime("%B %d, %Y")
        }

        doc = Document()

        # Title
        doc.add_heading('RESIDENTIAL LEASE AGREEMENT', 0)

        # Sections with placeholder content
        doc.add_paragraph(f'1. THE PARTIES. This Residential Lease Agreement (“Agreement”) is made on {data["agreement_date"]} by and between:')
        doc.add_paragraph(f'Landlord\nName: {data["landlord_name"]}\nMailing Address: {data["landlord_address"]}\n')
        doc.add_paragraph(f'Tenant\nName: {data["tenant_name"]}\nMailing Address: {data["tenant_address"]}\n')
        doc.add_paragraph(f'2. PROPERTY\nProperty Address: {data["property_address"]}\nResidence Type: {data["residence_type"]}\nBedroom(s): {data["bedrooms"]} Bathroom(s): {data["bathrooms"]}\n')
        doc.add_paragraph(f'3. TERM\nStart Date: {data["start_date"]}\nEnd Date: {data["end_date"]}\n')
        doc.add_paragraph(f'4. RENT\nMonthly Rent: ${data["rent_amount"]}\nDue Date: 1st day of each month.\n')
        doc.add_paragraph(f'5. SECURITY DEPOSIT\nAmount: ${data["security_deposit"]}\n')
        doc.add_paragraph('IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first above written.\n')

        # Add signature lines
        doc.add_paragraph("Landlord's Signature: __________________________")
        doc.add_paragraph("Tenant's Signature: __________________________")

        # Add signatures above the lines
        add_signature_above(doc, data["landlord_name"], output_dir, fonts_folder, 'Landlord')
        add_signature_above(doc, data["tenant_name"], output_dir, fonts_folder, 'Tenant')

        # Save the document
        file_name = f'Residential_Lease_Agreement_{data["tenant_name"].replace(" ", "_")}_{data["landlord_name"].replace(" ", "_")}.docx'
        doc.save(os.path.join(output_dir, file_name))

        print(f"Lease Agreement created for {data['landlord_name']} and {data['tenant_name']}.")

def create_nda(num_docs=1, output_dir="NDA_Documents", fonts_folder="fonts"):
    os.makedirs(output_dir, exist_ok=True)

    for _ in range(num_docs):
        data = {
            "disclosing_party_name": fake.company(),
            "disclosing_party_address": fake.address().replace("\n", ", "),
            "receiving_party_name": fake.name(),
            "receiving_party_address": fake.address().replace("\n", ", "),
            "agreement_date": fake.date_between(start_date="-1y", end_date="today").strftime("%B %d, %Y")
        }

        doc = Document()

        # Title
        doc.add_heading('NON-DISCLOSURE AGREEMENT (NDA)', 0)

        # Introduction
        doc.add_paragraph(f'This Nondisclosure Agreement ("Agreement") is made on {data["agreement_date"]} by and between the following parties:\n')
        doc.add_paragraph(f'Disclosing Party: {data["disclosing_party_name"]}\nAddress: {data["disclosing_party_address"]}\n')
        doc.add_paragraph(f'Receiving Party: {data["receiving_party_name"]}\nAddress: {data["receiving_party_address"]}\n')

        # Agreement Sections
        sections = {
            "1. Definition of Confidential Information": "Confidential Information includes all information or material that has, or could have, commercial value.",
            "2. Exclusions from Confidential Information": "Information that is publicly known at the time of disclosure is not considered confidential.",
            "3. Obligations of Receiving Party": "The Receiving Party agrees to maintain confidentiality and not disclose the information.",
            "4. Time Periods": "This Agreement remains in effect until the information is no longer confidential or upon mutual consent."
        }

        for title, content in sections.items():
            p = doc.add_paragraph()
            p.add_run(f"{title}: ").bold = True
            p.add_run(content)

        doc.add_paragraph("\nIN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.\n")

        # Add signature lines
        doc.add_paragraph("Disclosing Party's Signature: __________________________")
        doc.add_paragraph("Receiving Party's Signature: __________________________")

        # Add signatures
        add_signature_above(doc, data["disclosing_party_name"], output_dir, fonts_folder, 'Disclosing Party')
        add_signature_above(doc, data["receiving_party_name"], output_dir, fonts_folder, 'Receiving Party')

        # Save the document
        file_name = f'NDA_Agreement_{data["disclosing_party_name"].replace(" ", "_")}_{data["receiving_party_name"].replace(" ", "_")}.docx'
        doc.save(os.path.join(output_dir, file_name))

        print(f"NDA created for {data['disclosing_party_name']} and {data['receiving_party_name']}.")

def create_partnership_agreement(num_docs=1, output_dir="Partnership_Agreements", fonts_folder="fonts"):
    os.makedirs(output_dir, exist_ok=True)

    for _ in range(num_docs):
        # Generate fake data
        data = {
            "partner1_name": fake.name(),
            "partner1_address": fake.street_address(),
            "partner1_city_state_zip": f"{fake.city()}, {fake.state_abbr()} {fake.zipcode()}",
            "partner2_name": fake.name(),
            "partner2_address": fake.street_address(),
            "partner2_city_state_zip": f"{fake.city()}, {fake.state_abbr()} {fake.zipcode()}",
            "business_purpose": fake.catch_phrase(),
            "business_name": fake.company(),
            "agreement_date": fake.date_between(start_date="-1y", end_date="today").strftime("%B %d, %Y"),
            "state_of_governance": fake.state()
        }

        doc = Document()

        # Title
        doc.add_heading('PARTNERSHIP AGREEMENT', 0)

        # Introduction
        doc.add_paragraph(
            f'This PARTNERSHIP AGREEMENT ("Agreement") is made this {data["agreement_date"]} '
            f'by and between the following individuals:\n'
        )

        # Partner Information
        doc.add_paragraph(
            f'Partner 1: {data["partner1_name"]}\n'
            f'Address: {data["partner1_address"]}, {data["partner1_city_state_zip"]}\n'
        )
        doc.add_paragraph(
            f'Partner 2: {data["partner2_name"]}\n'
            f'Address: {data["partner2_address"]}, {data["partner2_city_state_zip"]}\n'
        )

        # Sections of the Agreement
        sections = {
            "1. Nature of Business": (
                "The partners listed above agree to form a partnership for the purpose of conducting the following business: "
                f"{data['business_purpose']}."
            ),
            "2. Name of the Partnership": (
                f"The partnership shall operate under the name '{data['business_name']}' and maintain offices at [ADDRESS]."
            ),
            "3. Day-To-Day Operations": (
                "The partners shall devote their full time and best efforts to the business. "
                "All partners will have equal management rights unless otherwise agreed upon."
            ),
            "4. Capital Contributions": (
                f"Each partner agrees to contribute the following capital to the partnership:\n"
                f" - {data['partner1_name']}: $10,000 (50% Share)\n"
                f" - {data['partner2_name']}: $10,000 (50% Share)"
            ),
            "5. Profits and Losses": (
                "Profits and losses shall be divided according to the capital contributions of each partner unless otherwise agreed."
            ),
            "6. Termination": (
                "The partnership may be terminated upon mutual agreement. Upon termination, assets will be distributed based on each partner's share."
            ),
            "7. Governing Law": (
                f"This Agreement shall be governed by the laws of the state of {data['state_of_governance']}."
            )
        }

        # Add each section to the document
        for title, content in sections.items():
            p = doc.add_paragraph()
            p.add_run(f"{title}: ").bold = True
            p.add_run(content)

        # Signature Section
        doc.add_paragraph("\nIN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.\n")

        # Add signature lines
        doc.add_paragraph("Partner 1's Signature: __________________________")
        doc.add_paragraph("Partner 2's Signature: __________________________")

        # Add signatures above the lines
        add_signature_above(doc, data["partner1_name"], output_dir, fonts_folder, "Partner 1")
        add_signature_above(doc, data["partner2_name"], output_dir, fonts_folder, "Partner 2")

        # Save the document
        file_name = f'Partnership_Agreement_{data["partner1_name"].replace(" ", "_")}_{data["partner2_name"].replace(" ", "_")}.docx'
        doc.save(os.path.join(output_dir, file_name))

        print(f"Partnership Agreement created for {data['partner1_name']} and {data['partner2_name']}.")

def create_petition_for_dissolution(num_docs=1, output_dir="Dissolution_Petitions", fonts_folder="fonts"):
    os.makedirs(output_dir, exist_ok=True)

    for _ in range(num_docs):
        data = {
            "husband_name": fake.name_male(),
            "wife_name": fake.name_female(),
            "husband_address": fake.address().replace("\n", ", "),
            "wife_address": fake.address().replace("\n", ", "),
            "marriage_date": fake.date_between(start_date="-10y", end_date="-5y").strftime("%B %d, %Y"),
            "marriage_city": fake.city(),
            "marriage_state": fake.state(),
            "case_number": fake.random_int(min=100000, max=999999),
            "division": "Family",
            "filing_date": fake.date_between(start_date="-1y", end_date="today").strftime("%B %d, %Y"),
            "court_address": fake.address().replace("\n", ", "),
            "court_phone_number": fake.phone_number()
        }

        doc = Document()

        # Court Information
        doc.add_paragraph(
            'IN THE CIRCUIT COURT OF THE [Judicial Circuit] JUDICIAL CIRCUIT, '
            f'IN AND FOR [County] COUNTY, [State]\n'
            f'Case No.: {data["case_number"]}\nDivision: {data["division"]}\n'
        )
        doc.add_paragraph(f'{data["court_address"]}\nCourt telephone No.: {data["court_phone_number"]}\n')

        # Parties Information
        parties_paragraph = doc.add_paragraph()
        parties_paragraph.add_run(f'Husband: {data["husband_name"]}  ').bold = True
        parties_paragraph.add_run('and  ').bold = True
        parties_paragraph.add_run(f'Wife: {data["wife_name"]}\n').bold = True

        # Title
        doc.add_heading('PETITION FOR SIMPLIFIED DISSOLUTION OF MARRIAGE', 0)

        # Petition Body
        doc.add_paragraph(f'We, {data["husband_name"]}, Husband, and {data["wife_name"]}, Wife, being sworn, certify that the following information is true:\n')
        petition_sections = [
            '1. We are both asking the Court for a dissolution of our marriage.',
            f'2. Husband lives at {data["husband_address"]}. Wife lives at {data["wife_address"]}.',
            f'3. We were married on {data["marriage_date"]} in the city of {data["marriage_city"]}, {data["marriage_state"]}.',
            '4. Our marriage is irretrievably broken.',
            '5. We do not have any minor or dependent children together, and the wife is not pregnant.',
            '6. We have divided our assets and liabilities by agreement, and we are satisfied with the division.'
        ]

        for section in petition_sections:
            doc.add_paragraph(section)

        # Agreement Options
        doc.add_paragraph('[Check one only]')
        doc.add_paragraph(
            '- ( ) Our marital settlement agreement is attached and was signed freely by both parties.\n'
            '- ( ) We prefer to keep our financial agreements private.\n'
        )

        doc.add_paragraph(f'Filed on: {data["filing_date"]}\n')

        # Add signature lines
        doc.add_paragraph("Husband's Signature: __________________________")
        doc.add_paragraph("Wife's Signature: __________________________")

        # Add signatures
        add_signature_above(doc, data["husband_name"], output_dir, fonts_folder, "Husband")
        add_signature_above(doc, data["wife_name"], output_dir, fonts_folder, "Wife")

        # Save the document
        file_name = f'Petition_for_Simplified_Dissolution_{data["husband_name"].replace(" ", "_")}_{data["wife_name"].replace(" ", "_")}.docx'
        doc.save(os.path.join(output_dir, file_name))

        print(f"Petition for Simplified Dissolution created for {data['husband_name']} and {data['wife_name']}.")

# --- Main Execution ---

if __name__ == "__main__":
    create_nda(30)
    create_lease_agreement(30)
    create_partnership_agreement(30)
    create_petition_for_dissolution(30)
